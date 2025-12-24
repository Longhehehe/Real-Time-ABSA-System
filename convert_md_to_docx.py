import markdown
from docx import Document
from docx.shared import Pt, RGBColor
import re

def markdown_to_docx(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()
    
    # Simple parser for the markdown structure
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('!['):
            # Image handling: ![alt](path)
            match = re.search(r'\!\[.*\]\((.*?)\)', line)
            if match:
                img_path = match.group(1)
                try:
                    # Resize image to fit page width (approx 6 inches)
                    doc.add_picture(img_path, width=Pt(400))
                except Exception as e:
                    print(f"Warning: Could not add image {img_path}. Error: {e}")
                    doc.add_paragraph(f"[Image: {img_path}]")
        elif line.startswith('```'):
            # Code block marker, skip or handle
            continue
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    # Use the specific artifact path
    md_path = r"C:\Users\Long\.gemini\antigravity\brain\4df68bf3-e77a-45a3-a710-31f1233dadd0\REPORT_PROJECT.md"
    docx_path = r"c:/Users/Long/Documents/Hoc_Tap/SE363 (1)/REPORT_PROJECT.docx"
    
    try:
        markdown_to_docx(md_path, docx_path)
    except Exception as e:
        print(f"Error: {e}")
